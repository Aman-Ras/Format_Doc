from flask import Flask, request, render_template_string, send_file, flash
from dotenv import load_dotenv
import io
import os
import re
import json
import requests
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import PyPDF2
try:
    from pyngrok import ngrok
except Exception:
    ngrok = None

load_dotenv()
api_key = os.getenv('OPENAI_API_KEY')

app = Flask(__name__)
app.secret_key = 'groq_resume_formatter_2025'

# Network configuration for local network access
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

# Clean HTML Template
UPLOAD_FORM = '''
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>HTML-Based Resume Formatter</title>
    <style>
        body { 
            font-family: Arial, sans-serif; 
            background-color: #f5f5f5; 
            margin: 0; 
            padding: 20px; 
        }
        .container { 
            background: white; 
            max-width: 600px; 
            margin: 0 auto; 
            padding: 40px; 
            border-radius: 12px; 
            box-shadow: 0 4px 12px rgba(0,0,0,0.1); 
        }
        h1 { 
            text-align: center; 
            color: #2c3e50; 
            margin-bottom: 10px; 
            font-size: 28px; 
            font-weight: bold; 
        }
        .subtitle { 
            text-align: center; 
            color: #7f8c8d; 
            margin-bottom: 30px; 
            font-size: 14px; 
        }
        .upload-section { 
            border: 2px dashed #3498db; 
            padding: 30px; 
            margin: 25px 0; 
            background-color: #f8f9fa; 
            border-radius: 8px; 
        }
        .upload-section h3 { 
            margin: 0 0 20px 0; 
            color: #2c3e50; 
            font-size: 18px; 
        }
        .file-input { 
            margin: 15px 0; 
        }
        .file-input label { 
            display: block; 
            margin-bottom: 5px; 
            color: #2c3e50; 
            font-weight: bold; 
        }
        input[type="file"] { 
            padding: 8px; 
            border: 1px solid #ddd; 
            border-radius: 4px; 
            width: 100%; 
            box-sizing: border-box; 
        }
        .submit-btn { 
            background: #3498db; 
            color: white; 
            padding: 15px 40px; 
            border: none; 
            border-radius: 6px; 
            cursor: pointer; 
            font-size: 16px; 
            font-weight: bold; 
            display: block; 
            margin: 0 auto; 
            transition: background-color 0.3s; 
        }
        .submit-btn:hover { 
            background: #2980b9; 
        }
        .file-info { 
            font-size: 12px; 
            color: #7f8c8d; 
            margin-top: 5px; 
        }
        .error { 
            color: #e74c3c; 
            margin: 15px 0; 
            padding: 15px; 
            background: #fdf2f2; 
            border-radius: 6px; 
            border: 1px solid #fecaca; 
        }
    .loading-overlay {
        position: fixed;
        top: 0;
        left: 0;
        right: 0;
        bottom: 0;
        background: rgba(255, 255, 255, 0.85);
        display: none;
        align-items: center;
        justify-content: center;
        z-index: 9999;
    }
    .loading-box {
        display: flex;
        flex-direction: column;
        align-items: center;
        gap: 12px;
        background: #ffffff;
        padding: 24px 28px;
        border-radius: 10px;
        box-shadow: 0 6px 18px rgba(0,0,0,0.12);
        border: 1px solid #e5e7eb;
    }
    .spinner {
        width: 42px;
        height: 42px;
        border: 4px solid #3498db;
        border-top-color: transparent;
        border-radius: 50%;
        animation: spin 1s linear infinite;
    }
    @keyframes spin {
        to { transform: rotate(360deg); }
    }
    .loading-text {
        color: #2c3e50;
        font-weight: bold;
        font-size: 14px;
    }
    </style>
</head>
<body>
    <div class="container">
        <h1>HTML-Based Resume Formatter</h1>
        <p class="subtitle">Perfect format matching from extracted HTML</p>
        
        {% with messages = get_flashed_messages() %}
            {% if messages %}
                <div class="error">
                    {% for message in messages %}
                        <p><strong>Error:</strong> {{ message }}</p>
                    {% endfor %}
                </div>
            {% endif %}
        {% endwith %}
        
        <form id="uploadForm" method="POST" enctype="multipart/form-data">
            <div class="upload-section">
                <h3>Upload Files</h3>
                
                <div class="file-input">
                    <label for="resume_file">Resume:</label>
                    <input type="file" id="resume_file" name="resume_file" accept=".pdf,.doc,.docx,.txt" required>
                </div>
                
                <div class="file-input">
                    <label for="logo_file">Logo (Optional):</label>
                    <input type="file" id="logo_file" name="logo_file" accept=".png,.jpg,.jpeg">
                    <div class="file-info">Will be named logo.png</div>
                </div>
            </div>
            
            <button id="submitBtn" type="submit" class="submit-btn">Create Professional Resume</button>
        </form>
        
        <div id="loadingOverlay" class="loading-overlay">
            <div class="loading-box">
                <div class="spinner"></div>
                <div class="loading-text">Formatting your resume...</div>
            </div>
        </div>
    </div>
<script>
  (function() {
    var form = document.getElementById('uploadForm');
    if (!form) return;
    var overlay = document.getElementById('loadingOverlay');
    var btn = document.getElementById('submitBtn');

    function showOverlay() {
      if (btn) {
        btn.disabled = true;
        btn.textContent = 'Generating...';
      }
      if (overlay) {
        overlay.style.display = 'flex';
      }
    }

    function hideOverlay() {
      if (overlay) {
        overlay.style.display = 'none';
      }
      if (btn) {
        btn.disabled = false;
        btn.textContent = 'Create Professional Resume';
      }
    }

    form.addEventListener('submit', function(e) {
      e.preventDefault();
      showOverlay();

      var formData = new FormData(form);
      fetch('/', {
        method: 'POST',
        body: formData
      }).then(function(response) {
        var disposition = response.headers.get('Content-Disposition') || '';
        var contentType = (response.headers.get('Content-Type') || '').toLowerCase();

        if (disposition.indexOf('attachment') !== -1 || contentType.indexOf('application/vnd.openxmlformats-officedocument.wordprocessingml.document') !== -1) {
          var filename = 'formatted_resume.docx';
          var match = /filename\*=UTF-8''([^;]+)|filename="?([^";]+)"?/i.exec(disposition);
          if (match) {
            filename = decodeURIComponent(match[1] || match[2] || filename);
          }
          return response.blob().then(function(blob) {
            var url = window.URL.createObjectURL(blob);
            var a = document.createElement('a');
            a.href = url;
            a.download = filename;
            document.body.appendChild(a);
            a.click();
            a.remove();
            window.URL.revokeObjectURL(url);
            hideOverlay();
          });
        }

        return response.text().then(function(html) {
          document.open();
          document.write(html);
          document.close();
        });
      }).catch(function() {
        hideOverlay();
        alert('Network error while generating the resume. Please try again.');
      });
    });
  })();
</script>
</body>
</html>
'''

# Helper Functions for DOCX Formatting
def set_cell_background_white(cell):
    """Set table cell background to white"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'FFFFFF')
    cell._element.get_or_add_tcPr().append(shading)

def set_calibri_font(run, size=11, bold=False, color=RGBColor(0, 0, 0)):
    """Set Calibri font with consistent styling"""
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

def remove_table_borders(table):
    """Remove table borders for clean look"""
    try:
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = tcPr.first_child_found_in("w:tcBorders")
                if tcBorders is not None:
                    tcPr.remove(tcBorders)
    except:
        pass

# OpenAI API Resume Extractor Class
class OpenAIResumeExtractor:
    def __init__(self, api_key=None):
        self.api_key = api_key or os.getenv('OPENAI_API_KEY')
        self.api_url = "https://api.openai.com/v1/chat/completions"
        # Default to a fast JSON-capable model; adjust as needed
        self.model = "gpt-4.1-mini"
    
    def _clean_json_text(self, text):
        """Clean model text to maximize chances of valid JSON parsing."""
        if not isinstance(text, str):
            return "{}"
        cleaned = text.strip()
        # Strip fenced code blocks ```json ... ``` or ``` ... ```
        if cleaned.startswith("```"):
            cleaned = re.sub(r"^```(?:json)?\n", "", cleaned, flags=re.IGNORECASE)
            cleaned = re.sub(r"\n```$", "", cleaned)
        # Extract content between the first '{' and the last '}' if present
        json_start = cleaned.find('{')
        json_end = cleaned.rfind('}')
        if json_start != -1 and json_end != -1 and json_end > json_start:
            cleaned = cleaned[json_start:json_end+1]
        # Normalize smart quotes
        cleaned = cleaned.replace('\u201c', '"').replace('\u201d', '"').replace('\u2019', "'")
        cleaned = cleaned.replace('“', '"').replace('”', '"').replace('’', "'")
        # Remove trailing commas before closing braces/brackets
        cleaned = re.sub(r",\s*(\}|\])", r"\1", cleaned)
        return cleaned
    
    def extract_with_openai(self, resume_text):
        """Extract structured data using OpenAI API"""
        if not self.api_key:
            print("OPENAI API key not provided, falling back to rule-based extraction")
            return self.extract_with_rules(resume_text)
        
        # Enhanced prompt for comprehensive extraction
        prompt = f"""
        You are a professional resume parser. Extract structured information from this resume text and return ONLY valid JSON.

        Resume Text:
        {resume_text}

        Extract and return a JSON object with these exact fields:
        {{
            "name": "Full candidate name",
            "location": "City/Location mentioned",
            "email": "Email address if found, otherwise empty string",
            "phone": "Phone number if found, otherwise empty string", 
            "date": "Application date if found, otherwise current date in DD-MMM-YYYY format",
            "subject": "Subject line in the exact format 'Application for the Position of {{ROLE}}' (no leading 'Subject:' label)",
            "summary": "Professional summary/objective from Key Expertise section and first paragraph (4-5 lines)",
            "education": [
                {{"degree": "Degree name", "institution": "Institution name"}}
            ],
            "experience_table": [
                {{
                    "company_name": "Company Name as Role (Duration)",
                    "roles_responsibility": [
                        "Exact responsibility 1 as written in resume",
                        "Exact responsibility 2 as written in resume",
                        "Exact responsibility 3 as written in resume"
                    ]
                }}
            ],
            "skills": ["List of technical skills, tools, technologies from Key Skills section"],
            "certifications": ["List of certifications and qualifications"],
            "cover_letter": "Complete cover letter content including all paragraphs after 'Dear Hiring Manager,'. If no cover letter text exists in the resume, GENERATE a professional, concise cover letter tailored to the 'subject' and the candidate's profile, using proper paragraph structure (2-5 paragraphs)."
        }}
        
        CRITICAL RULES:
        1. Extract ONLY information clearly present in the text
        2. For "name": Extract the candidate's full name as shown in the document header
        3. For "location": Extract city/location (e.g., "Bengaluru", "Bangalore")
        4. For "date": Extract application date if found, otherwise use current date
        5. For "subject": Return exactly "Application for the Position of {{ROLE}}" using the candidate's target role/title; do not include a leading "Subject:" label or extra punctuation
        6. For "summary": Combine Key Expertise section and first paragraph of cover letter
        7. For "education": Extract degree and institution in table format
        8. For "experience_table": Format each entry to match the table structure with "Company Name as Role (Duration)" in company_name field
        9. For "roles_responsibility" field: Extract each responsibility EXACTLY as written by the candidate from the Roles & Responsibility column
        10. For "skills": Extract from Key Skills section as a bulleted list
        11. For "certifications": Extract from CERTIFICATIONS section
        12. For "cover_letter": GENERATE a professional cover letter using the candidate's details (summary, key skills, certifications, experience highlights) tailored to the subject, and return it as multi-paragraph text (avoid repeating "Dear Hiring Manager,").
        13. Preserve the original formatting, punctuation, and exact language used by the candidate
        14. Return ONLY valid JSON with proper array formatting
        15. Ensure all JSON syntax is correct with proper quotes and commas
        """

        
        headers = {
            'Authorization': f'Bearer {self.api_key}',
            'Content-Type': 'application/json'
        }
        
        data = {
            'model': self.model,
            'messages': [
                {'role': 'system', 'content': 'You are a professional resume parsing expert. Extract information accurately and return only valid JSON.'},
                {'role': 'user', 'content': prompt}
            ],
            'temperature': 0.1,
            'max_tokens': 2500,
            'top_p': 1,
            'stream': False,
            # Enforce pure JSON response when supported
            'response_format': {'type': 'json_object'}
        }
        
        try:
            response = requests.post(self.api_url, headers=headers, json=data, timeout=120)
            response.raise_for_status()
            
            result = response.json()
            extracted_text = result['choices'][0]['message']['content'].strip()
            # Clean and parse JSON response
            json_str = self._clean_json_text(extracted_text)
            try:
                extracted_data = json.loads(json_str)
                return self.validate_and_clean_data(extracted_data)
            except json.JSONDecodeError as inner_e:
                print(f"JSON parsing error after cleaning: {str(inner_e)}")
                print("Raw model output (truncated):", extracted_text[:1000])
                return self.extract_with_rules(resume_text)
                
        except requests.exceptions.RequestException as e:
            print(f"OpenAI API request error: {str(e)}")
            return self.extract_with_rules(resume_text)
        except json.JSONDecodeError as e:
            print(f"JSON parsing error: {str(e)}")
            return self.extract_with_rules(resume_text)
        except Exception as e:
            print(f"OpenAI extraction error: {str(e)}")
            return self.extract_with_rules(resume_text)
    
    def validate_and_clean_data(self, data):
        """Validate and clean extracted data"""
        defaults = {
            'name': '',
            'location': '',
            'email': '',
            'phone': '',
            'date': '',
            'subject': '',
            'summary': '',
            'education': [],
            'experience_table': [],
            'experience': [],  # Keep for backward compatibility
            'skills': [],
            'certifications': [],
            'cover_letter': ''
        }
        
        # Convert education format if needed
        if 'education' in data and isinstance(data['education'], list):
            education_tuples = []
            for edu in data['education']:
                if isinstance(edu, dict):
                    degree = edu.get('degree', 'BCA')
                    institution = edu.get('institution', 'University Name')
                    education_tuples.append((degree, institution))
                elif isinstance(edu, (list, tuple)) and len(edu) >= 2:
                    education_tuples.append((edu[0], edu[1]))
            data['education'] = education_tuples if education_tuples else defaults['education']
        
        # Merge with defaults for missing fields
        for key, default_value in defaults.items():
            if key not in data or not data[key]:
                data[key] = default_value
        
        # Clean and limit arrays
        if isinstance(data.get('skills'), list):
            data['skills'] = [skill.strip() for skill in data['skills'][:15] if isinstance(skill, str) and skill.strip()]
        
        if isinstance(data.get('certifications'), list):
            data['certifications'] = [cert.strip() for cert in data['certifications'][:10] if isinstance(cert, str) and cert.strip()]

        
        # Use LLM-provided cover letter as-is (trim only)
        if isinstance(data.get('cover_letter'), str):
            data['cover_letter'] = data['cover_letter'].strip()
        else:
            data['cover_letter'] = ''
        
        # Ensure cover letter ends with desired closing
        if data['cover_letter']:
            candidate_name = (data.get('name') or '').strip() or 'Candidate'
            desired_closing = f"Sincerely,\n{candidate_name}"
            current = data['cover_letter'].rstrip()
            if not current.endswith(desired_closing):
                data['cover_letter'] = current + "\n\n" + desired_closing

        return data

    def extract_with_rules(self, resume_text):
        """Enhanced rule-based extraction as fallback"""
        lines = [line.strip() for line in resume_text.split('\n') if line.strip()]
        
        info = {
            'name': 'Professional Candidate',
            'location': 'Bengaluru',
            'email': '',
            'phone': '',
            'date': datetime.now().strftime("%d-%b-%Y"),
            'subject': 'Application for Technical Position',
            'summary': '',
            'education': [('BCA', 'University Name')],
            'experience': [],
            'skills': [],
            'certifications': []
        }
        
        # Extract name (first meaningful line)
        for line in lines[:20]:
            if (len(line.split()) <= 4 and len(line) > 2 and 
                not any(char.isdigit() for char in line[:15]) and 
                '@' not in line and 
                not any(keyword in line.lower() for keyword in ['over', 'years', 'experience', 'seeking'])):
                info['name'] = line
                break
        
        # Extract contact info
        email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', resume_text)
        if email_match:
            info['email'] = email_match.group()
        
        phone_match = re.search(r'[+]?[\d\s\-\(\)]{10,15}', resume_text)
        if phone_match:
            info['phone'] = phone_match.group().strip()
        
        # Extract location
        location_keywords = ['bengaluru', 'bangalore', 'mumbai', 'delhi', 'chennai', 'hyderabad', 'pune']
        for line in lines:
            if any(city in line.lower() for city in location_keywords) and len(line) < 50:
                info['location'] = line.title()
                break
        
        # Extract summary (first substantial paragraph)
        for line in lines:
            if (len(line) > 100 and ('years' in line.lower() or 'experience' in line.lower())):
                info['summary'] = line[:300] + '...' if len(line) > 300 else line
                break
        
        # Extract skills
        skill_keywords = ['AWS', 'Azure', 'Linux', 'Windows', 'Docker', 'Kubernetes', 'Python', 'Java', 
                         'Terraform', 'Ansible', 'Jenkins', 'Git', 'SQL', 'Oracle', 'MySQL', 'MongoDB',
                         'Windchill', 'PLM', 'Cisco', 'RHEL', 'CentOS', 'Ubuntu']
        skills_found = set()
        
        for line in lines:
            for skill in skill_keywords:
                if skill.lower() in line.lower():
                    skills_found.add(skill)
        
        info['skills'] = list(skills_found)[:12]
        
        # Extract certifications
        for line in lines:
            if ('certified' in line.lower() or 'certification' in line.lower()) and len(line.strip()) > 10:
                cleaned_cert = line.strip().replace('•', '').replace('-', '').strip()
                if cleaned_cert:
                    info['certifications'].append(cleaned_cert)
        
        return info

# Professional Resume Formatter Class
class ProfessionalResumeFormatter:
    def __init__(self, logo_path=None):
        self.doc = Document()
        self.logo_path = logo_path
        self.blue_color = RGBColor(0, 51, 102)    # #003366 - for "Key Skills & Technologies" only
        self.black_color = RGBColor(0, 0, 0)      # #000000 - for all other text
        self.setup_document()
    
    def setup_document(self):
        """Set document margins"""
        for section in self.doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
    
    def add_header_with_logo(self, name, location):
        """Add header with logo and name (RIGHT aligned)"""
        header_table = self.doc.add_table(rows=1, cols=2)
        
        # Logo cell
        logo_cell = header_table.cell(0, 0)
        if self.logo_path and os.path.exists(self.logo_path):
            try:
                para = logo_cell.paragraphs[0]
                para.clear()
                run = para.add_run()
                run.add_picture(self.logo_path, width=Inches(2.5))
            except:
                logo_cell.text = "⊞ Logo"
        else:
            logo_cell.text = "⊞ Logo"
        
        # Name cell - RIGHT aligned (ONLY name, no location)
        name_cell = header_table.cell(0, 1)
        name_para = name_cell.paragraphs[0]
        name_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        name_run = name_para.add_run(name)
        set_calibri_font(name_run, 18, True, self.black_color)
        
        # Location removed from header - will be added separately in document body
        
        remove_table_borders(header_table)
        self.doc.add_paragraph()
    
    def add_page_header(self, name):
        """Add header for new pages (logo + name only)"""
        header_table = self.doc.add_table(rows=1, cols=2)
        
        # Logo cell
        logo_cell = header_table.cell(0, 0)
        if self.logo_path and os.path.exists(self.logo_path):
            try:
                para = logo_cell.paragraphs[0]
                para.clear()
                run = para.add_run()
                run.add_picture(self.logo_path, width=Inches(2.5))
            except:
                logo_cell.text = "⊞ Logo"
        else:
            logo_cell.text = "⊞ Logo"
        
        # Name cell - RIGHT aligned
        name_cell = header_table.cell(0, 1)
        name_para = name_cell.paragraphs[0]
        name_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        name_run = name_para.add_run(name)
        set_calibri_font(name_run, 18, True, self.black_color)
        
        remove_table_borders(header_table)
        self.doc.add_paragraph()
    
    def add_objective(self, text):
        """Add justified objective paragraph"""
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        set_calibri_font(run, 11, False, self.black_color)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.space_after = Pt(4)  # Reduced from 12 to 6 points
    
    def add_section_header(self, text, use_blue=False):
        """Add section headers"""
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        color = self.blue_color if use_blue else self.black_color
        set_calibri_font(run, 12, True, color)
        para.space_after = Pt(0)  # No spacing after section headers
    
    def add_blue_header(self, text):
        """Add blue background header like 'Key Expertise' and 'Education Details'"""
        # Create a table with 1 row and 1 column for the blue background
        header_table = self.doc.add_table(rows=1, cols=1)
        header_table.allow_autofit = False
        header_table.autofit = False
        
        # Set table width to use full page width
        header_table.columns[0].width = Inches(6.5)
        
        # Get the cell and set blue background
        cell = header_table.rows[0].cells[0]
        cell.text = text
        
        # Set darker light blue background
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), 'DCE6F0')  # Darker light blue color
        cell._element.get_or_add_tcPr().append(shading)
        
        # Format text (black, bold, left-aligned)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                set_calibri_font(run, 12, True, RGBColor(0, 0, 0))  # Black text
        
        # Remove table borders
        remove_table_borders(header_table)
        
        # Add minimal spacing after header
        self.doc.add_paragraph()
        # Reduce spacing after blue header
        last_para = self.doc.paragraphs[-1]
        last_para.space_after = Pt(3)  # Reduced spacing
    
    def add_page_header_to_all_pages(self, name):
        """Add header with logo and name that appears on every page"""
        # Get the section
        section = self.doc.sections[0]
        
        # Create header
        header = section.header
        header.paragraphs.clear()  # Clear any existing header content
        
        # Create header table with logo on left and name on right
        header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
        header_table.allow_autofit = False
        header_table.autofit = False
        
        # Set column widths - reduced gaps
        header_table.columns[0].width = Inches(3.0)  # Logo column - reduced from 3.5
        header_table.columns[1].width = Inches(3.5)  # Name column - increased from 3.0
        
        # Left cell - Logo
        logo_cell = header_table.rows[0].cells[0]
        logo_para = logo_cell.paragraphs[0]
        logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add logo image if available
        if self.logo_path and os.path.exists(self.logo_path):
            try:
                logo_run = logo_para.add_run()
                logo_run.add_picture(self.logo_path, width=Inches(2.5))
            except Exception as e:
                print(f"Error adding logo to header: {e}")
                # Fallback to text if image fails
                logo_run = logo_para.add_run("⊞ Logo")
                set_calibri_font(logo_run, 10, False, self.black_color)
        else:
            # Fallback to text if no logo file
            logo_run = logo_para.add_run("⊞ Logo")
            set_calibri_font(logo_run, 10, False, self.black_color)
        
        # Right cell - Name
        name_cell = header_table.rows[0].cells[1]
        name_para = name_cell.paragraphs[0]
        name_run = name_para.add_run(name)
        set_calibri_font(name_run, 16, True, self.black_color)  # Increased from 12 to 16
        name_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Remove table borders
        remove_table_borders(header_table)
    
    def add_table(self, headers, data_rows):
        """Add table with labels on left, content on right"""
        try:
            print(f"Starting table creation...")
            print(f"Data rows: {data_rows}")
            print(f"Number of data rows: {len(data_rows)}")
            
            # Create table with 2 columns (labels on left, content on right)
            table = self.doc.add_table(rows=len(data_rows), cols=2)
            print(f"Table created successfully")
            
            # Set table style
            table.style = 'Table Grid'
            print(f"Table style set")
            
            # Set table properties
            table.autofit = False
            table.allow_autofit = False
            
            # Set column widths - labels on left, content on right
            table.columns[0].width = Inches(2.0)  # Labels column (Company Name, Roles & Responsibility)
            table.columns[1].width = Inches(5.0)  # Content column (actual company info and responsibilities)
            print(f"Column widths set")
            
            # Data rows - no header row needed
            for row_idx, row_data in enumerate(data_rows):
                print(f"Processing row {row_idx}: {row_data}")
                data_row = table.rows[row_idx]
                
                # Left cell - Label (Company Name, Roles & Responsibility)
                left_cell = data_row.cells[0]
                left_cell.text = row_data[0]  # Label text
                set_cell_background_white(left_cell)
                
                # Format left cell (labels) - bold text
                for paragraph in left_cell.paragraphs:
                    # Center-align "Roles & Responsibility", left-align "Company Name"
                    if row_data[0] == "Roles & Responsibility":
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # Force center alignment by setting paragraph properties
                        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        set_calibri_font(run, 11, True, self.black_color)  # Bold for labels
                
                # Right cell - Content (company info or responsibilities)
                right_cell = data_row.cells[1]
                right_cell.text = ""
                
                # Handle multi-line text with bullet points
                if isinstance(row_data[1], str) and '\n' in row_data[1]:
                    lines = row_data[1].split('\n')
                    for i, line in enumerate(lines):
                        if line.strip():
                            if i == 0:
                                para = right_cell.paragraphs[0]
                                para.text = line.strip()
                            else:
                                para = right_cell.add_paragraph()
                                para.text = line.strip()
                else:
                    right_cell.text = str(row_data[1])
                
                # Format right cell (content)
                set_cell_background_white(right_cell)
                for paragraph in right_cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        # Make company name bold, responsibilities regular
                        if row_data[0] == "Company Name":  # If this is the company name row
                            set_calibri_font(run, 11, True, self.black_color)  # Bold for company name
                        else:
                            set_calibri_font(run, 11, False, self.black_color)  # Regular text for responsibilities
                
                print(f"Row {row_idx} processed")
            
            # Add spacing after table
            self.doc.add_paragraph()
            print(f"Table creation completed successfully")
            
        except Exception as e:
            print(f"Error creating table: {str(e)}")
            import traceback
            traceback.print_exc()
            # Fallback: add simple text instead of table
            self.doc.add_paragraph("Error creating table - displaying as text:")
            for row_data in data_rows:
                self.doc.add_paragraph(f"{row_data[0]}: {row_data[1]}")
                self.doc.add_paragraph()
    
    def add_standard_table(self, headers, data_rows):
        """Add standard table with headers at top (for education, etc.)"""
        try:
            print(f"Starting standard table creation...")
            print(f"Headers: {headers}")
            print(f"Data rows: {data_rows}")
            print(f"Number of data rows: {len(data_rows)}")
            
            # Create table with header row + data rows
            table = self.doc.add_table(rows=len(data_rows) + 1, cols=len(headers))
            print(f"Table created successfully")
            
            # Set table style
            table.style = 'Table Grid'
            print(f"Table style set")
            
            # Set table properties
            table.autofit = False
            table.allow_autofit = False
            
            # Set column widths
            table.columns[0].width = Inches(3.0)  # Qualification column
            table.columns[1].width = Inches(4.0)  # Institution column
            print(f"Column widths set")
            
            # Header row
            header_row = table.rows[0]
            for i, header in enumerate(headers):
                cell = header_row.cells[i]
                cell.text = header
                set_cell_background_white(cell)
                
                # Format header text - bold
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        set_calibri_font(run, 11, True, self.black_color)
            print(f"Header row created")
            
            # Data rows
            for row_idx, row_data in enumerate(data_rows, 1):
                print(f"Processing row {row_idx}: {row_data}")
                data_row = table.rows[row_idx]
                for col_idx, cell_text in enumerate(row_data):
                    cell = data_row.cells[col_idx]
                    cell.text = str(cell_text)
                    
                    # Format cell
                    set_cell_background_white(cell)
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        for run in paragraph.runs:
                            set_calibri_font(run, 11, False, self.black_color)
                print(f"Row {row_idx} processed")
            
            # Add spacing after table
            self.doc.add_paragraph()
            print(f"Standard table creation completed successfully")
            
        except Exception as e:
            print(f"Error creating standard table: {str(e)}")
            import traceback
            traceback.print_exc()
            # Fallback: add simple text instead of table
            self.doc.add_paragraph("Error creating standard table - displaying as text:")
            for row_data in data_rows:
                self.doc.add_paragraph(f"{headers[0]}: {row_data[0]}, {headers[1]}: {row_data[1]}")
                self.doc.add_paragraph()
    
    def add_bullet_list(self, items):
        """Add bullet point list with minimal spacing"""
        for item in items:
            para = self.doc.add_paragraph()
            run = para.add_run(f"• {item}")
            set_calibri_font(run, 11, False, self.black_color)
            # Eliminate all spacing between bullet points
            para.space_after = Pt(0)  # No spacing after paragraph
            para.space_before = Pt(0)  # No spacing before paragraph
            # Set line spacing to single
            para.paragraph_format.line_spacing = 1.0
    
    def create_complete_resume(self, candidate_info):
        """Create complete professional resume"""
        # Add page header that appears on every page
        self.add_page_header_to_all_pages(candidate_info['name'])
        
        # Objective/Summary (only if present)
        if candidate_info.get('summary'):
            self.add_blue_header("Key Expertise")
            self.add_objective(candidate_info['summary'])
        
        # Education table (only if present)
        if candidate_info.get('education'):
            self.add_blue_header("Education Details")
            self.add_standard_table(["Qualification", "Institution"], candidate_info['education'])
        
        # Contact info - compact spacing
        name_para = self.doc.add_paragraph()
        name_run = name_para.add_run(candidate_info['name'])
        set_calibri_font(name_run, 16, True, self.black_color)
        
        if candidate_info.get('location'):
            loc_para = self.doc.add_paragraph()
            loc_run = loc_para.add_run(candidate_info['location'])
            set_calibri_font(loc_run, 11, False, self.black_color)
            loc_para.space_after = Pt(3)  # Minimal spacing after location
        
        # Date and Subject - compact spacing
        if candidate_info.get('date'):
            date_para = self.doc.add_paragraph()
            date_run = date_para.add_run(f"Date: {candidate_info['date']}")
            set_calibri_font(date_run, 11, False, self.black_color)
            date_para.space_after = Pt(3)  # Minimal spacing after date
        
        if candidate_info.get('subject'):
            subj_para = self.doc.add_paragraph()
            subj_label = subj_para.add_run("Subject: ")
            set_calibri_font(subj_label, 11, True, self.black_color)
            subj_text = subj_para.add_run(str(candidate_info['subject']))
            set_calibri_font(subj_text, 11, False, self.black_color)
            subj_para.space_after = Pt(3)  # Minimal spacing after subject
        
        # Cover letter (only if present)
        if candidate_info.get('cover_letter'):
            dear_para = self.doc.add_paragraph()
            dear_run = dear_para.add_run("Dear Hiring Manager,")
            set_calibri_font(dear_run, 11, True, self.black_color)  # Made bold
            # Split cover letter into paragraphs and add each separately
            cover_paragraphs = candidate_info['cover_letter'].split('\n\n')
            for para_text in cover_paragraphs:
                if para_text.strip() and "Dear Hiring Manager" not in para_text.strip():
                    cover_para = self.doc.add_paragraph()
                    cover_run = cover_para.add_run(para_text.strip())
                    set_calibri_font(cover_run, 11, False, self.black_color)
                    cover_para.space_after = Pt(3)  # Reduced spacing
        
        # Skills section - Ultra compact spacing like certifications
        if candidate_info.get('skills'):
            self.add_section_header("Key Skills")
            # Add skills with ultra-compact spacing
            for item in candidate_info['skills']:
                para = self.doc.add_paragraph()
                run = para.add_run(f"\t• {item}")
                set_calibri_font(run, 11, False, self.black_color)
                # Ultra-compact spacing for skills
                para.space_after = Pt(0)  # No spacing after
                para.space_before = Pt(0)  # No spacing before
                para.paragraph_format.line_spacing = 1.0  # Single line spacing
                para.paragraph_format.space_after = Pt(0)  # Additional Word-level spacing control
            # Insert a single blank line after skills list
            self.doc.add_paragraph()
 
        # Experience section
        if candidate_info.get('experience_table'):
            self.add_section_header("Experience:")
            print(f"Found {len(candidate_info['experience_table'])} experience entries")
            print(f"Experience table data: {candidate_info['experience_table']}")
            
            for exp in candidate_info['experience_table']:
                print(f"Processing experience entry: {exp}")
                
                # Convert array of responsibilities to bullet points
                responsibilities_text = ""
                if isinstance(exp['roles_responsibility'], list):
                    for i, resp in enumerate(exp['roles_responsibility'], 1):
                        responsibilities_text += f"• {resp}\n"
                    responsibilities_text = responsibilities_text.strip()
                else:
                    responsibilities_text = str(exp['roles_responsibility'])
                
                print(f"Company: {exp['company_name']}")
                print(f"Responsibilities: {responsibilities_text[:100]}...")
                
                # Create table with labels on left, content on right
                exp_data = [
                    ("Company Name", exp['company_name']),
                    ("Roles & Responsibility", responsibilities_text)
                ]
                print(f"Creating table with data: {exp_data}")
                self.add_table(["", ""], exp_data)  # Empty headers since labels are in left column
                # Spacing after table is already added inside add_table(); avoid extra spacing
        elif candidate_info.get('experience'):  # Fallback for old format
            self.add_section_header("Experience:")
            print(f"Using fallback experience format: {candidate_info['experience']}")
            for exp in candidate_info['experience'][:3]:
                # Convert array of responsibilities to bullet points
                responsibilities_text = ""
                if isinstance(exp['responsibilities'], list):
                    for i, resp in enumerate(exp['responsibilities'], 1):
                        responsibilities_text += f"• {resp}\n"
                    responsibilities_text = responsibilities_text.strip()
                else:
                    responsibilities_text = str(exp['responsibilities'])
                
                # Create table with labels on left, content on right
                exp_data = [
                    ("Company Name", f"{exp['company']} as {exp['role']} ({exp['duration']})"),
                    ("Roles & Responsibility", responsibilities_text)
                ]
                self.add_table(["", ""], exp_data)  # Empty headers since labels are in left column
                # Spacing after table is already added inside add_table(); avoid extra spacing
        
        # Certifications - ultra compact spacing
        if candidate_info.get('certifications'):
            self.add_section_header("CERTIFICATIONS")
            # Add certifications with ultra-compact spacing
            for item in candidate_info['certifications']:
                para = self.doc.add_paragraph()
                run = para.add_run(f"• {item}")
                set_calibri_font(run, 11, False, self.black_color)
                # Ultra-compact spacing for certifications
                para.space_after = Pt(0)  # No spacing after
                para.space_before = Pt(0)  # No spacing before
                para.paragraph_format.line_spacing = 1.0  # Single line spacing
                para.paragraph_format.space_after = Pt(0)  # Additional spacing control
        
        # Debug: Check if we have any experience data at all
        if not candidate_info.get('experience_table') and not candidate_info.get('experience'):
            print("WARNING: No experience data found at all!")
            print(f"All available keys: {list(candidate_info.keys())}")
            self.doc.add_paragraph("No experience data available")
        
        return self.save_to_buffer()
    
    def save_to_buffer(self):
        """Save to memory buffer"""
        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer

# Text Extraction Classes
class ResumeTextExtractor:
    def extract_from_pdf(self, file_content):
        try:
            reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text is None:
                    page_text = ""
                text += page_text + "\n"
            return text
        except Exception as e:
            return f"PDF Error: {str(e)}"
    
    def extract_from_docx(self, file_content):
        try:
            doc = Document(io.BytesIO(file_content))
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
                
            # ✅ NEW: Extract text from tables also
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data:
                        text += " | ".join(row_data) + "\n"
                        
            return text
        except Exception as e:
            return f"DOCX Error: {str(e)}"
    
    def extract_from_txt(self, file_content):
        try:
            return file_content.decode('utf-8', errors='ignore')
        except Exception as e:
            return f"TXT Error: {str(e)}"

def extract_text_from_file(file):
    """Extract text from uploaded file"""
    extractor = ResumeTextExtractor()
    content = file.read()
    filename = file.filename.lower()
    
    if filename.endswith('.pdf'):
        return extractor.extract_from_pdf(content)
    elif filename.endswith(('.doc', '.docx')):
        return extractor.extract_from_docx(content)
    elif filename.endswith('.txt'):
        return extractor.extract_from_txt(content)
    else:
        return "Unsupported file format"

# Main Flask Route
@app.route('/', methods=['GET', 'POST'])
def upload_and_format():
    if request.method == 'POST':
        # Validate resume file
        if 'resume_file' not in request.files or request.files['resume_file'].filename == '':
            flash('Please select a resume file')
            return render_template_string(UPLOAD_FORM)
        
        resume_file = request.files['resume_file']
        
        # Handle optional logo - use default if none provided
        logo_path = "logo_path.png"  # Default logo path
        if 'logo_file' in request.files and request.files['logo_file'].filename != '':
            logo_file = request.files['logo_file']
            logo_path = "temp_logo.png"
            logo_file.save(logo_path)
        
        try:
            # Extract text from resume
            resume_text = extract_text_from_file(resume_file)
            
            # Print the raw extracted text from the uploaded document
            print("\n" + "="*80)
            print("RAW EXTRACTED RESUME TEXT")
            print("="*80)
            print(resume_text)
            print("="*80 + "\n")
            
            if resume_text.startswith(('PDF Error', 'DOCX Error', 'TXT Error', 'Unsupported')):
                flash(f'Error processing file: {resume_text}')
                return render_template_string(UPLOAD_FORM)
            
            if len(resume_text.strip()) < 50:
                flash('Resume appears to be empty or too short')
                return render_template_string(UPLOAD_FORM)
            
            # OpenAI extraction (set OPENAI_API_KEY environment variable)
            api_key = os.getenv('OPENAI_API_KEY')
            extractor = OpenAIResumeExtractor(api_key)
            candidate_info = extractor.extract_with_openai(resume_text)
            
            # Debug: Print extracted data
            print("\n" + "="*80)
            print("LLM EXTRACTION RESULTS")
            print("="*80)
            print(f"Name: {candidate_info.get('name', 'NOT FOUND')}")
            print(f"Location: {candidate_info.get('location', 'NOT FOUND')}")
            print(f"Date: {candidate_info.get('date', 'NOT FOUND')}")
            print(f"Subject: {candidate_info.get('subject', 'NOT FOUND')}")
            print(f"Email: {candidate_info.get('email', 'NOT FOUND')}")
            print(f"Phone: {candidate_info.get('phone', 'NOT FOUND')}")
            print(f"Summary: {candidate_info.get('summary', 'NOT FOUND')[:200]}...")
            print(f"Education: {candidate_info.get('education', 'NOT FOUND')}")
            print(f"Skills: {candidate_info.get('skills', 'NOT FOUND')}")
            print(f"Certifications: {candidate_info.get('certifications', 'NOT FOUND')}")
            print(f"Cover Letter: {candidate_info.get('cover_letter', 'NOT FOUND')[:200]}...")
            
            print("\n" + "-"*80)
            print("EXPERIENCE TABLE DETAILS")
            print("-"*80)
            if candidate_info.get('experience_table'):
                for i, exp in enumerate(candidate_info['experience_table'], 1):
                    print(f"\nExperience {i}:")
                    print(f"  Company Name: {exp.get('company_name', 'NOT FOUND')}")
                    print(f"  Roles & Responsibility: {exp.get('roles_responsibility', 'NOT FOUND')}")
            else:
                print("No experience_table found!")
                
            print("\n" + "-"*80)
            print("FALLBACK EXPERIENCE (OLD FORMAT)")
            print("-"*80)
            if candidate_info.get('experience'):
                for i, exp in enumerate(candidate_info['experience'], 1):
                    print(f"\nExperience {i}:")
                    print(f"  Company: {exp.get('company', 'NOT FOUND')}")
                    print(f"  Role: {exp.get('role', 'NOT FOUND')}")
                    print(f"  Duration: {exp.get('duration', 'NOT FOUND')}")
                    print(f"  Responsibilities: {exp.get('responsibilities', 'NOT FOUND')}")
            else:
                print("No experience (old format) found!")
                
            print("\n" + "="*80)
            print("COMPLETE JSON RESPONSE")
            print("="*80)
            import json
            print(json.dumps(candidate_info, indent=2, ensure_ascii=False))
            print("="*80 + "\n")
            
            # Generate professional resume
            formatter = ProfessionalResumeFormatter(logo_path)
            buffer = formatter.create_complete_resume(candidate_info)
            
            # Cleanup temporary logo (only if it's a temp file)
            if logo_path == "temp_logo.png" and os.path.exists(logo_path):
                os.remove(logo_path)
            
            # Return formatted resume
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            safe_name = re.sub(r'[^\w\s-]', '', candidate_info['name']).strip().replace(' ', '_')
            filename = f"{safe_name}_{timestamp}.docx"
            
            return send_file(
                buffer,
                as_attachment=True,
                download_name=filename,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            
        except Exception as e:
            # Cleanup on error (only temp logo)
            if 'logo_path' in locals() and logo_path == "temp_logo.png" and os.path.exists(logo_path):
                os.remove(logo_path)
            flash(f'Processing error: {str(e)}')
            return render_template_string(UPLOAD_FORM)
    
    return render_template_string(UPLOAD_FORM)

if __name__ == '__main__':
    import socket
    
    # Get local IP address for network access
    def get_local_ip():
        try:
            # Connect to a remote address to get local IP
            s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
            s.connect(("8.8.8.8", 80))
            local_ip = s.getsockname()[0]
            s.close()
            return local_ip
        except:
            return "127.0.0.1"
    
    local_ip = get_local_ip()
    port = 5005
    
    print("=" * 60)
    print("🌐 RESUME FORMATTER - NETWORK ACCESS")
    print("=" * 60)
    print(f"Local access: http://localhost:{port}")
    print(f"Network access: http://{local_ip}:{port}")
    print(f"Other users on same network can use: http://{local_ip}:{port}")
    print("=" * 60)
    print("Starting server...")

    # Start ngrok tunnel if pyngrok is installed
    if ngrok is not None:
        try:
            ngrok_authtoken = os.getenv('NGROK_AUTHTOKEN')
            if ngrok_authtoken:
                ngrok.set_auth_token(ngrok_authtoken)
            http_tunnel = ngrok.connect(addr=port, proto="http")
            public_url = http_tunnel.public_url
            print("=" * 60)
            print(f"Public (ngrok) URL: {public_url}")
            print("Share this URL to access your app over the internet.")
            print("=" * 60)
        except Exception as e:
            print(f"ngrok start failed: {e}")
    else:
        print("pyngrok not installed; you can run 'ngrok http {port}' separately to expose the app.")
    
    app.run(debug=True, host='0.0.0.0', port=port, use_reloader=False)
